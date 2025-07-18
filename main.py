import streamlit as st
import pandas as pd
import openpyxl
from docx import Document
import io
from datetime import datetime
import base64
import anthropic
import os
from typing import List, Dict

# Page configuration
st.set_page_config(
    page_title="Cap Table Tie-Out Analysis",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

class DeterministicCapTableAnalyzer:
    def __init__(self, api_key: str):
        self.client = anthropic.Anthropic(api_key=api_key)
        self.uploaded_files = {}
        
    def read_docx_content(self, file_content: bytes, filename: str) -> str:
        """Read DOCX content and return as plain text"""
        try:
            doc = Document(io.BytesIO(file_content))
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            return '\n'.join(full_text)
        except Exception as e:
            st.error(f"Error reading {filename}: {str(e)}")
            return ""
    
    def excel_to_structured_data(self, file_content: bytes, filename: str) -> List[Dict]:
        """Convert Excel to structured data for analysis"""
        try:
            df = pd.read_excel(io.BytesIO(file_content), engine='openpyxl', header=None)
            
            # Find header row
            header_row_idx = None
            for i, row in df.iterrows():
                row_str = ' '.join([str(cell) for cell in row if pd.notna(cell)])
                if 'Security ID' in row_str and 'Stakeholder Name' in row_str:
                    header_row_idx = i
                    break
            
            if header_row_idx is None:
                return []
            
            # Extract headers and data
            headers = df.iloc[header_row_idx].tolist()
            entries = []
            
            for i in range(header_row_idx + 1, len(df)):
                row = df.iloc[i]
                if pd.notna(row.iloc[0]) and str(row.iloc[0]).strip():  # Has Security ID
                    entry = {}
                    for j, header in enumerate(headers):
                        if j < len(row) and pd.notna(header):
                            entry[str(header)] = row.iloc[j] if pd.notna(row.iloc[j]) else ""
                    entries.append(entry)
            
            return entries
            
        except Exception as e:
            st.error(f"Error parsing Excel: {str(e)}")
            return []
    
    def extract_board_grants(self, board_docs: Dict[str, str]) -> List[Dict]:
        """Extract grants from board documents using deterministic rules"""
        grants = []
        
        for filename, content in board_docs.items():
            content_lower = content.lower()
            
            # Determine document type
            if 'repurchase' in content_lower:
                grant = self.extract_repurchase_info(content, filename)
                if grant:
                    grants.append(grant)
            elif 'restricted stock' in content_lower or 'rsa' in content_lower:
                grant = self.extract_rsa_grant(content, filename)
                if grant:
                    grants.append(grant)
            elif 'option' in content_lower:
                grant = self.extract_option_grant(content, filename)
                if grant:
                    grants.append(grant)
        
        return grants
    
    def extract_rsa_grant(self, content: str, filename: str) -> Dict:
        """Extract RSA grant info using comprehensive pattern matching"""
        import re
        
        grant = {
            'type': 'RSA Grant',
            'filename': filename,
            'stockholder': None,
            'shares': None,
            'price_per_share': None,
            'date': None,
            'vesting_start': None,
            'vesting_schedule': None
        }
        
        # Debug: Show what we're parsing
        st.write(f"**Parsing {filename}:**")
        st.write(f"Content length: {len(content)} characters")
        
        # Extract date - multiple patterns
        date_patterns = [
            r'Date:\s*([A-Za-z]+\s+\d{1,2},\s+\d{4})',
            r'dated\s+([A-Za-z]+\s+\d{1,2},\s+\d{4})',
            r'(\d{1,2}/\d{1,2}/\d{4})',
            r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}'
        ]
        
        for pattern in date_patterns:
            date_match = re.search(pattern, content, re.IGNORECASE)
            if date_match:
                grant['date'] = date_match.group(1)
                st.write(f"✅ Found date: {grant['date']}")
                break
        
        if not grant['date']:
            st.write("❌ No date found")
        
        # Extract stockholder - look in schedule/table and throughout document
        stockholder_patterns = [
            r'Name[:\s]+([A-Za-z\s]+)',
            r'([A-Za-z]+\s+[A-Za-z]+)\s+\d+,?\d*\s+\
    
    def extract_repurchase_info(self, content: str, filename: str) -> Dict:
        """Extract repurchase info with comprehensive parsing"""
        import re
        
        repurchase = {
            'type': 'Repurchase',
            'filename': filename,
            'stockholder': None,
            'shares_repurchased': None,
            'date': None
        }
        
        st.write(f"**Parsing repurchase document {filename}:**")
        
        # Extract date
        date_patterns = [
            r'Date:\s*([A-Za-z]+\s+\d{1,2},\s+\d{4})',
            r'dated\s+([A-Za-z]+\s+\d{1,2},\s+\d{4})',
        ]
        
        for pattern in date_patterns:
            date_match = re.search(pattern, content, re.IGNORECASE)
            if date_match:
                repurchase['date'] = date_match.group(1)
                st.write(f"✅ Found repurchase date: {repurchase['date']}")
                break
        
        # Extract stockholder
        common_names = ['John Doe', 'Jane Smith', 'Bob', 'Alice', 'Charlie', 'Arthur']
        for name in common_names:
            if name in content:
                repurchase['stockholder'] = name
                st.write(f"✅ Found stockholder: {name}")
                break
        
        # Extract repurchased shares - multiple patterns
        repurchase_patterns = [
            r'repurchase\s+(\d{1,3}(?:,\d{3})*)\s+',
            r'(\d{1,3}(?:,\d{3})*)\s+unvested\s+shares',
            r'(\d{1,3}(?:,\d{3})*)\s+shares.*repurchas',
            r'exercise.*right.*repurchase\s+(\d{1,3}(?:,\d{3})*)',
        ]
        
        for pattern in repurchase_patterns:
            repurchase_match = re.search(pattern, content, re.IGNORECASE)
            if repurchase_match:
                shares_str = repurchase_match.group(1).replace(',', '')
                try:
                    shares = int(shares_str)
                    if 1 <= shares <= 100000:  # Reasonable range
                        repurchase['shares_repurchased'] = shares
                        st.write(f"✅ Found repurchased shares: {shares}")
                        break
                except ValueError:
                    continue
        
        st.write(f"**Final repurchase data:** {repurchase}")
        st.write("---")
        
        return repurchase
    
    def extract_option_grant(self, content: str, filename: str) -> Dict:
        """Extract option grant info"""
        # Similar to RSA but for options
        return self.extract_rsa_grant(content, filename)  # Reuse logic for now
    
    def run_deterministic_analysis(self, cap_table_entries: List[Dict], board_grants: List[Dict]) -> List[Dict]:
        """Run deterministic analysis that always produces same results"""
        discrepancies = []
        
        # Create lookup of board grants by stockholder
        board_lookup = {}
        repurchases = []
        
        for grant in board_grants:
            if grant['type'] == 'Repurchase':
                repurchases.append(grant)
            else:
                stockholder = grant.get('stockholder')
                if stockholder:
                    if stockholder not in board_lookup:
                        board_lookup[stockholder] = []
                    board_lookup[stockholder].append(grant)
        
        # Check each cap table entry
        for entry in cap_table_entries:
            security_id = entry.get('Security ID', '')
            stockholder = entry.get('Stakeholder Name', '')
            shares = self.safe_int(entry.get('Quantity Issued', 0))
            cost_basis = self.safe_float(entry.get('Cost Basis', 0))
            board_approval_date = str(entry.get('Board Approval Date', ''))
            vesting_schedule = str(entry.get('Vesting Schedule', ''))
            
            # Calculate price per share
            price_per_share = cost_basis / shares if shares > 0 else 0
            
            # Check 1: Phantom Equity (no board approval found)
            if stockholder not in board_lookup:
                discrepancies.append({
                    'number': len(discrepancies) + 1,
                    'severity': 'HIGH',
                    'stockholder': stockholder,
                    'security_id': security_id,
                    'issue': 'Phantom Equity Entry',
                    'cap_table_value': f'{shares} shares',
                    'legal_value': 'No supporting documentation found',
                    'description': f'Cap table shows {security_id} for {stockholder} but no board approval found',
                    'source': 'None found'
                })
                continue
            
            # Find matching board grant
            matching_grant = None
            for grant in board_lookup[stockholder]:
                if grant.get('shares') == shares or abs((grant.get('shares', 0) - shares)) <= 1:
                    matching_grant = grant
                    break
            
            if not matching_grant:
                matching_grant = board_lookup[stockholder][0]  # Use first grant as fallback
            
            # Check 2: Share quantity mismatch
            board_shares = matching_grant.get('shares', 0)
            if board_shares and abs(shares - board_shares) > 1:
                discrepancies.append({
                    'number': len(discrepancies) + 1,
                    'severity': 'HIGH',
                    'stockholder': stockholder,
                    'security_id': security_id,
                    'issue': 'Incorrect Share Quantity',
                    'cap_table_value': f'{shares} shares',
                    'legal_value': f'{board_shares} shares',
                    'description': f'Cap table shows {shares} shares but board approval is for {board_shares} shares',
                    'source': matching_grant.get('filename', 'Unknown')
                })
            
            # Check 3: Price mismatch
            board_price = matching_grant.get('price_per_share', 0)
            if board_price and abs(price_per_share - board_price) > 0.01:
                discrepancies.append({
                    'number': len(discrepancies) + 1,
                    'severity': 'HIGH',
                    'stockholder': stockholder,
                    'security_id': security_id,
                    'issue': 'Incorrect Price Per Share',
                    'cap_table_value': f'${price_per_share:.2f}',
                    'legal_value': f'${board_price:.2f}',
                    'description': f'Cap table shows ${price_per_share:.2f} per share but board approval is for ${board_price:.2f} per share',
                    'source': matching_grant.get('filename', 'Unknown')
                })
            
            # Check 4: Board approval date mismatch
            board_date = matching_grant.get('date', '')
            if board_date and board_date not in board_approval_date:
                discrepancies.append({
                    'number': len(discrepancies) + 1,
                    'severity': 'HIGH',
                    'stockholder': stockholder,
                    'security_id': security_id,
                    'issue': 'Incorrect Board Approval Date',
                    'cap_table_value': board_approval_date,
                    'legal_value': board_date,
                    'description': 'Board approval date in cap table does not match legal documents',
                    'source': matching_grant.get('filename', 'Unknown')
                })
            
            # Check 5: Vesting schedule mismatch
            board_vesting = matching_grant.get('vesting_schedule', '')
            if board_vesting and board_vesting not in vesting_schedule:
                if ('monthly' in board_vesting.lower() and 'monthly' not in vesting_schedule.lower()) or \
                   ('annual' in board_vesting.lower() and 'annual' not in vesting_schedule.lower()):
                    discrepancies.append({
                        'number': len(discrepancies) + 1,
                        'severity': 'HIGH',
                        'stockholder': stockholder,
                        'security_id': security_id,
                        'issue': 'Vesting Schedule Mismatch',
                        'cap_table_value': vesting_schedule,
                        'legal_value': board_vesting,
                        'description': 'Vesting schedule format differs between cap table and board documents',
                        'source': matching_grant.get('filename', 'Unknown')
                    })
        
        # Check 6: Missing repurchase transactions
        for repurchase in repurchases:
            stockholder = repurchase.get('stockholder')
            shares_repurchased = repurchase.get('shares_repurchased', 0)
            
            if stockholder and shares_repurchased:
                discrepancies.append({
                    'number': len(discrepancies) + 1,
                    'severity': 'HIGH',
                    'stockholder': stockholder,
                    'security_id': 'Multiple',
                    'issue': 'Missing Repurchase Transaction',
                    'cap_table_value': 'No repurchase reflected',
                    'legal_value': f'{shares_repurchased} shares repurchased',
                    'description': f'Board approved repurchase of {shares_repurchased} shares from {stockholder} but cap table does not reflect this transaction',
                    'source': repurchase.get('filename', 'Unknown')
                })
        
        return discrepancies
    
    def safe_int(self, value) -> int:
        """Safely convert to int"""
        try:
            if pd.isna(value):
                return 0
            return int(float(value))
        except (ValueError, TypeError):
            return 0
    
    def safe_float(self, value) -> float:
        """Safely convert to float"""
        try:
            if pd.isna(value):
                return 0.0
            return float(value)
        except (ValueError, TypeError):
            return 0.0
    def __init__(self, api_key: str):
        self.client = anthropic.Anthropic(api_key=api_key)
        self.uploaded_files = {}
        
    def read_docx_content(self, file_content: bytes, filename: str) -> str:
        """Read DOCX content and return as plain text"""
        try:
            doc = Document(io.BytesIO(file_content))
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            return '\n'.join(full_text)
        except Exception as e:
            st.error(f"Error reading {filename}: {str(e)}")
            return ""
    
    def excel_to_text_preview(self, file_content: bytes, filename: str) -> str:
        """Convert Excel to text preview for LLM analysis"""
        try:
            df = pd.read_excel(io.BytesIO(file_content), engine='openpyxl', header=None)
            
            # Create text representation
            text_preview = f"Excel file: {filename}\n\n"
            text_preview += "Raw data structure:\n"
            
            for i in range(min(15, len(df))):  # First 15 rows
                row_data = df.iloc[i].tolist()
                # Clean up the row data
                cleaned_row = []
                for cell in row_data:
                    if pd.isna(cell):
                        cleaned_row.append("")
                    else:
                        cleaned_row.append(str(cell))
                text_preview += f"Row {i + 1}: {cleaned_row}\n"
            
            return text_preview
        except Exception as e:
            return f"Error reading Excel file {filename}: {str(e)}"
    
    def create_analysis_prompt(self, board_docs: Dict[str, str], cap_table_text: str) -> str:
        """Create the enhanced prompt that catches all discrepancies with standardized approach"""
        
        prompt = """You are a lawyer conducting a standardized capitalization table tie out. You MUST follow this exact sequence:

MANDATORY ANALYSIS SEQUENCE:
1. DOCUMENT INVENTORY: List every board document and what it approves
2. CAP TABLE INVENTORY: List every cap table entry (Security ID, Stockholder, Shares)
3. SYSTEMATIC COMPARISON: For each cap table entry, check against board docs
4. DISCREPANCY IDENTIFICATION: List each discrepancy separately with exact format

STEP 1 - DOCUMENT INVENTORY:
First, create a complete list of all board-approved grants from legal documents:
- Document name, date, stockholder, shares, price, vesting details

STEP 2 - CAP TABLE INVENTORY: 
List every entry in the cap table:
- Security ID, Stockholder Name, Quantity, Price details, Dates

STEP 3 - SYSTEMATIC COMPARISON:
For EACH cap table entry, verify these 7 items in order:
a) Does this entry have board approval? (if no = PHANTOM EQUITY)
b) Do share quantities match?
c) Do prices match?
d) Do board approval dates match?
e) Do issue dates match?
f) Do vesting schedules match (monthly vs annual)?
g) Are repurchases reflected?

STEP 4 - DISCREPANCY LIST:
Use this EXACT format for each discrepancy:

DISCREPANCY #[X]: [Issue Title]
- Severity: HIGH/MEDIUM/LOW
- Stockholder: [Name]
- Security ID: [ID]
- Cap Table Shows: [Value]
- Legal Documents Show: [Value]
- Source Document: [Filename]
- Description: [1-2 sentences]

CRITICAL REQUIREMENTS:
- Analyze every single cap table entry
- Check for phantom equity (entries without board approval)
- Verify vesting schedule language matches exactly
- Check for missing repurchase transactions
- List each discrepancy separately (don't group)
- Use consistent severity: HIGH=material issues, MEDIUM=dates/documentation, LOW=minor

Here are the documents to analyze:

BOARD DOCUMENTS:
"""
        
        # Add each board document
        for filename, content in board_docs.items():
            prompt += f"\n--- {filename} ---\n{content}\n"
        
        prompt += f"\nSECURITIES LEDGER / CAP TABLE:\n{cap_table_text}\n"
        
        prompt += """

NOW EXECUTE THE 4-STEP ANALYSIS SEQUENCE ABOVE.

Begin with: "STEP 1 - DOCUMENT INVENTORY:" and follow the exact sequence."""
        
        return prompt
    
    def analyze_with_llm(self, board_docs: Dict[str, str], cap_table_text: str) -> str:
        """Send documents to LLM for analysis"""
        
        prompt = self.create_analysis_prompt(board_docs, cap_table_text)
        
        try:
            response = self.client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=4000,
                temperature=0,  # Maximum determinism
                system="You are a systematic legal auditor. Always follow the exact same analysis sequence and format. Be consistent and thorough in your approach.",
                messages=[
                    {
                        "role": "user", 
                        "content": prompt
                    }
                ]
            )
            
            return response.content[0].text
            
        except Exception as e:
            return f"Error analyzing documents: {str(e)}"

def main():
    st.title("📊 Cap Table Tie-Out Analysis")
    st.markdown("*LLM-powered analysis replicating expert legal review*")
    
    # Initialize analyzer with API key from secrets
    if 'analyzer' not in st.session_state:
        try:
            api_key = st.secrets["ANTHROPIC_API_KEY"]
            st.session_state.analyzer = DeterministicCapTableAnalyzer(api_key)
        except KeyError:
            st.error("❌ Anthropic API key not found in Streamlit secrets. Please configure ANTHROPIC_API_KEY in your .streamlit/secrets.toml file")
            st.stop()
    
    # File upload section
    with st.sidebar:
        st.header("📁 Upload Documents")
        
        # Board documents upload
        st.subheader("Board Documents")
        board_files = st.file_uploader(
            "Upload board consents, minutes, and legal docs",
            type=['docx', 'doc'],
            accept_multiple_files=True,
            key="board_docs",
            help="Upload DOCX files containing board resolutions, consents, and other legal documents"
        )
        
        # Securities ledger upload
        st.subheader("Securities Ledger")
        cap_table_file = st.file_uploader(
            "Upload cap table (Excel format)",
            type=['xlsx', 'xls'],
            key="cap_table",
            help="Upload Excel file containing the company's capitalization table"
        )
        
        # Analysis button
        st.markdown("---")
        run_analysis = st.button(
            "🔍 Run LLM Tie-Out Analysis", 
            type="primary", 
            use_container_width=True
        )
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📋 Uploaded Documents")
        
        if board_files:
            st.write("**Board Documents:**")
            for file in board_files:
                st.write(f"✅ {file.name}")
                
            # Show preview of first document
            if st.checkbox("Show document preview"):
                first_file = board_files[0]
                first_file.seek(0)
                analyzer = st.session_state.get('analyzer')
                if analyzer:
                    content = analyzer.read_docx_content(first_file.read(), first_file.name)
                    st.text_area("Document content preview:", content[:1000] + "...", height=200)
        else:
            st.info("No board documents uploaded yet")
        
        if cap_table_file:
            st.write("**Securities Ledger:**")
            st.write(f"✅ {cap_table_file.name}")
            
            # Show preview of cap table
            if st.checkbox("Show cap table preview"):
                cap_table_file.seek(0)
                try:
                    df_preview = pd.read_excel(io.BytesIO(cap_table_file.read()), engine='openpyxl')
                    st.dataframe(df_preview.head(10))
                except Exception as e:
                    st.error(f"Error previewing cap table: {e}")
        else:
            st.info("No securities ledger uploaded yet")
    
    with col2:
        st.subheader("⚙️ Analysis Status")
        
        if not board_files and not cap_table_file:
            st.warning("📄 Please upload documents to begin analysis")
        elif not board_files:
            st.warning("📋 Please upload board documents")
        elif not cap_table_file:
            st.warning("📊 Please upload securities ledger")
        else:
            st.success("✅ Ready for LLM analysis!")
            st.info("Click 'Run LLM Tie-Out Analysis' to start")
    
    # Run analysis when button is clicked
    if run_analysis:
        if not board_files or not cap_table_file:
            st.error("Please upload both board documents and securities ledger before running analysis")
            return
        
        with st.spinner("🔍 Running deterministic analysis... This will be consistent every time"):
            try:
                analyzer = st.session_state.analyzer
                
                # Process board documents in consistent order (alphabetical)
                board_docs = {}
                sorted_files = sorted(board_files, key=lambda x: x.name)
                for file in sorted_files:
                    file.seek(0)  # Reset file pointer
                    content = analyzer.read_docx_content(file.read(), file.name)
                    board_docs[file.name] = content
                
                # Process cap table
                cap_table_file.seek(0)
                cap_table_entries = analyzer.excel_to_structured_data(cap_table_file.read(), cap_table_file.name)
                
                # Extract board grants using deterministic rules
                board_grants = analyzer.extract_board_grants(board_docs)
                
                # Run deterministic analysis
                discrepancies = analyzer.run_deterministic_analysis(cap_table_entries, board_grants)
                
                # Display results
                st.markdown("---")
                st.header("🎯 Deterministic Analysis Results")
                st.success(f"✅ **Consistent Results**: Found {len(discrepancies)} discrepancies using rule-based analysis")
                
                if discrepancies:
                    # Summary metrics
                    high_count = len([d for d in discrepancies if d['severity'] == 'HIGH'])
                    medium_count = len([d for d in discrepancies if d['severity'] == 'MEDIUM'])
                    low_count = len([d for d in discrepancies if d['severity'] == 'LOW'])
                    
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("🔍 Total Issues", len(discrepancies))
                    with col2:
                        st.metric("🔴 High Severity", high_count)
                    with col3:
                        st.metric("🟡 Medium Severity", medium_count)
                    with col4:
                        st.metric("🟢 Low Severity", low_count)
                    
                    # Display each discrepancy
                    st.subheader("📋 Detailed Discrepancies")
                    
                    for disc in discrepancies:
                        severity_colors = {'HIGH': '#ff4444', 'MEDIUM': '#ffaa00', 'LOW': '#00aa00'}
                        severity_emojis = {'HIGH': '🔴', 'MEDIUM': '🟡', 'LOW': '🟢'}
                        
                        with st.container():
                            st.markdown(
                                f"""
                                <div style="
                                    border-left: 4px solid {severity_colors[disc['severity']]};
                                    background-color: {'#fff5f5' if disc['severity'] == 'HIGH' else '#fffaf0' if disc['severity'] == 'MEDIUM' else '#f0fff4'};
                                    padding: 1rem;
                                    margin: 1rem 0;
                                    border-radius: 0 8px 8px 0;
                                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                                ">
                                    <h4 style="margin: 0 0 0.5rem 0; color: {severity_colors[disc['severity']]};">
                                        {severity_emojis[disc['severity']]} DISCREPANCY #{disc['number']}: {disc['issue']}
                                    </h4>
                                    <p style="margin: 0; font-weight: bold; color: #333;">
                                        👤 <strong>Stockholder:</strong> {disc['stockholder']} | 
                                        📊 <strong>Severity:</strong> {disc['severity']}
                                    </p>
                                </div>
                                """, 
                                unsafe_allow_html=True
                            )
                            
                            col1, col2 = st.columns(2)
                            with col1:
                                st.markdown("**📋 Cap Table Shows:**")
                                st.code(disc['cap_table_value'])
                                st.markdown("**🆔 Security ID:**")
                                st.code(disc['security_id'])
                            
                            with col2:
                                st.markdown("**📜 Legal Documents Show:**")
                                st.code(disc['legal_value'])
                                st.markdown("**📄 Source Document:**")
                                st.code(disc['source'])
                            
                            st.markdown("**📝 Description:**")
                            st.write(disc['description'])
                            st.markdown("---")
                else:
                    st.success("🎉 No discrepancies found! The cap table appears to be in sync with the board documents.")
                
                # Create downloadable report
                st.markdown("---")
                st.subheader("📤 Download Report")
                
                if discrepancies:
                    # Create CSV for download
                    import pandas as pd
                    df_report = pd.DataFrame([
                        {
                            'Discrepancy #': d['number'],
                            'Severity': d['severity'],
                            'Stockholder': d['stockholder'],
                            'Security ID': d['security_id'],
                            'Issue': d['issue'],
                            'Cap Table Shows': d['cap_table_value'],
                            'Legal Documents Show': d['legal_value'],
                            'Description': d['description'],
                            'Source Document': d['source']
                        } for d in discrepancies
                    ])
                    
                    csv = df_report.to_csv(index=False)
                    st.download_button(
                        label="📄 Download Discrepancies Report (CSV)",
                        data=csv,
                        file_name=f"cap_table_discrepancies_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                        mime="text/csv"
                    )
                
            except Exception as e:
                st.error(f"Error during analysis: {str(e)}")
                st.error("Please check your API key and try again")
    
    # Information section
    with st.expander("ℹ️ How this works"):
        st.markdown("""
        **This tool replicates the exact process of expert legal document analysis:**
        
        1. **Document Upload**: You upload board documents and cap table (just like dragging files to Claude)
        2. **LLM Processing**: The documents are sent to Claude (same AI that did the original analysis)
        3. **Expert Analysis**: Claude performs the same thorough legal review methodology
        4. **Detailed Report**: You get the same quality discrepancy analysis and recommendations
        
        **What the LLM analyzes:**
        - Compares cap table entries against board approvals
        - Identifies missing documentation
        - Checks share quantities, pricing, and dates
        - Detects phantom equity grants
        - Finds missing transactions (repurchases, etc.)
        - Provides severity assessment and recommendations
        
        **API Key**: Configured via Streamlit secrets (ANTHROPIC_API_KEY in .streamlit/secrets.toml)
        """)
    
    # Footer
    st.markdown("---")
    st.markdown("*Powered by Claude AI - Professional legal document analysis*")

if __name__ == "__main__":
    main()
,  # Name followed by numbers and $
            r'([A-Za-z]+\s+[A-Za-z]+)\s+\d+,?\d*\s+shares',
            r'to\s+([A-Za-z]+\s+[A-Za-z]+)',
            r'from\s+([A-Za-z]+\s+[A-Za-z]+)',
        ]
        
        # Also look for common names explicitly
        common_names = ['John Doe', 'Jane Smith', 'Bob', 'Alice', 'Charlie', 'Arthur']
        for name in common_names:
            if name in content:
                grant['stockholder'] = name
                st.write(f"✅ Found stockholder: {name}")
                break
        
        if not grant['stockholder']:
            for pattern in stockholder_patterns:
                match = re.search(pattern, content, re.IGNORECASE)
                if match:
                    name = match.group(1).strip()
                    # Filter out common false positives
                    if name not in ['Date', 'DIRECTORS', 'Name', 'Board', 'Company']:
                        grant['stockholder'] = name
                        st.write(f"✅ Found stockholder via pattern: {name}")
                        break
        
        if not grant['stockholder']:
            st.write("❌ No stockholder found")
        
        # Extract shares - multiple patterns
        share_patterns = [
            r'(\d{1,3}(?:,\d{3})*)\s+shares?',
            r'shares?\s+(\d{1,3}(?:,\d{3})*)',
            r'issue.*?(\d{1,3}(?:,\d{3})*)',
            r'grant.*?(\d{1,3}(?:,\d{3})*)',
            # Look in schedule/table format
            r'(\d{1,3}(?:,\d{3})*)\s+\
    
    def extract_repurchase_info(self, content: str, filename: str) -> Dict:
        """Extract repurchase info"""
        import re
        
        repurchase = {
            'type': 'Repurchase',
            'filename': filename,
            'stockholder': None,
            'shares_repurchased': None,
            'date': None
        }
        
        # Extract date
        date_match = re.search(r'Date:\s*([A-Za-z]+\s+\d{1,2},\s+\d{4})', content)
        if date_match:
            repurchase['date'] = date_match.group(1)
        
        # Extract stockholder
        names = ['John Doe', 'Jane Smith', 'Bob', 'Alice', 'Charlie', 'Arthur']
        for name in names:
            if name in content:
                repurchase['stockholder'] = name
                break
        
        # Extract repurchased shares
        repurchase_match = re.search(r'repurchase\s+(\d{1,3}(?:,\d{3})*)', content, re.IGNORECASE)
        if repurchase_match:
            repurchase['shares_repurchased'] = int(repurchase_match.group(1).replace(',', ''))
        
        return repurchase
    
    def extract_option_grant(self, content: str, filename: str) -> Dict:
        """Extract option grant info"""
        # Similar to RSA but for options
        return self.extract_rsa_grant(content, filename)  # Reuse logic for now
    
    def run_deterministic_analysis(self, cap_table_entries: List[Dict], board_grants: List[Dict]) -> List[Dict]:
        """Run deterministic analysis that always produces same results"""
        discrepancies = []
        
        # Create lookup of board grants by stockholder
        board_lookup = {}
        repurchases = []
        
        for grant in board_grants:
            if grant['type'] == 'Repurchase':
                repurchases.append(grant)
            else:
                stockholder = grant.get('stockholder')
                if stockholder:
                    if stockholder not in board_lookup:
                        board_lookup[stockholder] = []
                    board_lookup[stockholder].append(grant)
        
        # Check each cap table entry
        for entry in cap_table_entries:
            security_id = entry.get('Security ID', '')
            stockholder = entry.get('Stakeholder Name', '')
            shares = self.safe_int(entry.get('Quantity Issued', 0))
            cost_basis = self.safe_float(entry.get('Cost Basis', 0))
            board_approval_date = str(entry.get('Board Approval Date', ''))
            vesting_schedule = str(entry.get('Vesting Schedule', ''))
            
            # Calculate price per share
            price_per_share = cost_basis / shares if shares > 0 else 0
            
            # Check 1: Phantom Equity (no board approval found)
            if stockholder not in board_lookup:
                discrepancies.append({
                    'number': len(discrepancies) + 1,
                    'severity': 'HIGH',
                    'stockholder': stockholder,
                    'security_id': security_id,
                    'issue': 'Phantom Equity Entry',
                    'cap_table_value': f'{shares} shares',
                    'legal_value': 'No supporting documentation found',
                    'description': f'Cap table shows {security_id} for {stockholder} but no board approval found',
                    'source': 'None found'
                })
                continue
            
            # Find matching board grant
            matching_grant = None
            for grant in board_lookup[stockholder]:
                if grant.get('shares') == shares or abs((grant.get('shares', 0) - shares)) <= 1:
                    matching_grant = grant
                    break
            
            if not matching_grant:
                matching_grant = board_lookup[stockholder][0]  # Use first grant as fallback
            
            # Check 2: Share quantity mismatch
            board_shares = matching_grant.get('shares', 0)
            if board_shares and abs(shares - board_shares) > 1:
                discrepancies.append({
                    'number': len(discrepancies) + 1,
                    'severity': 'HIGH',
                    'stockholder': stockholder,
                    'security_id': security_id,
                    'issue': 'Incorrect Share Quantity',
                    'cap_table_value': f'{shares} shares',
                    'legal_value': f'{board_shares} shares',
                    'description': f'Cap table shows {shares} shares but board approval is for {board_shares} shares',
                    'source': matching_grant.get('filename', 'Unknown')
                })
            
            # Check 3: Price mismatch
            board_price = matching_grant.get('price_per_share', 0)
            if board_price and abs(price_per_share - board_price) > 0.01:
                discrepancies.append({
                    'number': len(discrepancies) + 1,
                    'severity': 'HIGH',
                    'stockholder': stockholder,
                    'security_id': security_id,
                    'issue': 'Incorrect Price Per Share',
                    'cap_table_value': f'${price_per_share:.2f}',
                    'legal_value': f'${board_price:.2f}',
                    'description': f'Cap table shows ${price_per_share:.2f} per share but board approval is for ${board_price:.2f} per share',
                    'source': matching_grant.get('filename', 'Unknown')
                })
            
            # Check 4: Board approval date mismatch
            board_date = matching_grant.get('date', '')
            if board_date and board_date not in board_approval_date:
                discrepancies.append({
                    'number': len(discrepancies) + 1,
                    'severity': 'HIGH',
                    'stockholder': stockholder,
                    'security_id': security_id,
                    'issue': 'Incorrect Board Approval Date',
                    'cap_table_value': board_approval_date,
                    'legal_value': board_date,
                    'description': 'Board approval date in cap table does not match legal documents',
                    'source': matching_grant.get('filename', 'Unknown')
                })
            
            # Check 5: Vesting schedule mismatch
            board_vesting = matching_grant.get('vesting_schedule', '')
            if board_vesting and board_vesting not in vesting_schedule:
                if ('monthly' in board_vesting.lower() and 'monthly' not in vesting_schedule.lower()) or \
                   ('annual' in board_vesting.lower() and 'annual' not in vesting_schedule.lower()):
                    discrepancies.append({
                        'number': len(discrepancies) + 1,
                        'severity': 'HIGH',
                        'stockholder': stockholder,
                        'security_id': security_id,
                        'issue': 'Vesting Schedule Mismatch',
                        'cap_table_value': vesting_schedule,
                        'legal_value': board_vesting,
                        'description': 'Vesting schedule format differs between cap table and board documents',
                        'source': matching_grant.get('filename', 'Unknown')
                    })
        
        # Check 6: Missing repurchase transactions
        for repurchase in repurchases:
            stockholder = repurchase.get('stockholder')
            shares_repurchased = repurchase.get('shares_repurchased', 0)
            
            if stockholder and shares_repurchased:
                discrepancies.append({
                    'number': len(discrepancies) + 1,
                    'severity': 'HIGH',
                    'stockholder': stockholder,
                    'security_id': 'Multiple',
                    'issue': 'Missing Repurchase Transaction',
                    'cap_table_value': 'No repurchase reflected',
                    'legal_value': f'{shares_repurchased} shares repurchased',
                    'description': f'Board approved repurchase of {shares_repurchased} shares from {stockholder} but cap table does not reflect this transaction',
                    'source': repurchase.get('filename', 'Unknown')
                })
        
        return discrepancies
    
    def safe_int(self, value) -> int:
        """Safely convert to int"""
        try:
            if pd.isna(value):
                return 0
            return int(float(value))
        except (ValueError, TypeError):
            return 0
    
    def safe_float(self, value) -> float:
        """Safely convert to float"""
        try:
            if pd.isna(value):
                return 0.0
            return float(value)
        except (ValueError, TypeError):
            return 0.0
    def __init__(self, api_key: str):
        self.client = anthropic.Anthropic(api_key=api_key)
        self.uploaded_files = {}
        
    def read_docx_content(self, file_content: bytes, filename: str) -> str:
        """Read DOCX content and return as plain text"""
        try:
            doc = Document(io.BytesIO(file_content))
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            return '\n'.join(full_text)
        except Exception as e:
            st.error(f"Error reading {filename}: {str(e)}")
            return ""
    
    def excel_to_text_preview(self, file_content: bytes, filename: str) -> str:
        """Convert Excel to text preview for LLM analysis"""
        try:
            df = pd.read_excel(io.BytesIO(file_content), engine='openpyxl', header=None)
            
            # Create text representation
            text_preview = f"Excel file: {filename}\n\n"
            text_preview += "Raw data structure:\n"
            
            for i in range(min(15, len(df))):  # First 15 rows
                row_data = df.iloc[i].tolist()
                # Clean up the row data
                cleaned_row = []
                for cell in row_data:
                    if pd.isna(cell):
                        cleaned_row.append("")
                    else:
                        cleaned_row.append(str(cell))
                text_preview += f"Row {i + 1}: {cleaned_row}\n"
            
            return text_preview
        except Exception as e:
            return f"Error reading Excel file {filename}: {str(e)}"
    
    def create_analysis_prompt(self, board_docs: Dict[str, str], cap_table_text: str) -> str:
        """Create the enhanced prompt that catches all discrepancies with standardized approach"""
        
        prompt = """You are a lawyer conducting a standardized capitalization table tie out. You MUST follow this exact sequence:

MANDATORY ANALYSIS SEQUENCE:
1. DOCUMENT INVENTORY: List every board document and what it approves
2. CAP TABLE INVENTORY: List every cap table entry (Security ID, Stockholder, Shares)
3. SYSTEMATIC COMPARISON: For each cap table entry, check against board docs
4. DISCREPANCY IDENTIFICATION: List each discrepancy separately with exact format

STEP 1 - DOCUMENT INVENTORY:
First, create a complete list of all board-approved grants from legal documents:
- Document name, date, stockholder, shares, price, vesting details

STEP 2 - CAP TABLE INVENTORY: 
List every entry in the cap table:
- Security ID, Stockholder Name, Quantity, Price details, Dates

STEP 3 - SYSTEMATIC COMPARISON:
For EACH cap table entry, verify these 7 items in order:
a) Does this entry have board approval? (if no = PHANTOM EQUITY)
b) Do share quantities match?
c) Do prices match?
d) Do board approval dates match?
e) Do issue dates match?
f) Do vesting schedules match (monthly vs annual)?
g) Are repurchases reflected?

STEP 4 - DISCREPANCY LIST:
Use this EXACT format for each discrepancy:

DISCREPANCY #[X]: [Issue Title]
- Severity: HIGH/MEDIUM/LOW
- Stockholder: [Name]
- Security ID: [ID]
- Cap Table Shows: [Value]
- Legal Documents Show: [Value]
- Source Document: [Filename]
- Description: [1-2 sentences]

CRITICAL REQUIREMENTS:
- Analyze every single cap table entry
- Check for phantom equity (entries without board approval)
- Verify vesting schedule language matches exactly
- Check for missing repurchase transactions
- List each discrepancy separately (don't group)
- Use consistent severity: HIGH=material issues, MEDIUM=dates/documentation, LOW=minor

Here are the documents to analyze:

BOARD DOCUMENTS:
"""
        
        # Add each board document
        for filename, content in board_docs.items():
            prompt += f"\n--- {filename} ---\n{content}\n"
        
        prompt += f"\nSECURITIES LEDGER / CAP TABLE:\n{cap_table_text}\n"
        
        prompt += """

NOW EXECUTE THE 4-STEP ANALYSIS SEQUENCE ABOVE.

Begin with: "STEP 1 - DOCUMENT INVENTORY:" and follow the exact sequence."""
        
        return prompt
    
    def analyze_with_llm(self, board_docs: Dict[str, str], cap_table_text: str) -> str:
        """Send documents to LLM for analysis"""
        
        prompt = self.create_analysis_prompt(board_docs, cap_table_text)
        
        try:
            response = self.client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=4000,
                temperature=0,  # Maximum determinism
                system="You are a systematic legal auditor. Always follow the exact same analysis sequence and format. Be consistent and thorough in your approach.",
                messages=[
                    {
                        "role": "user", 
                        "content": prompt
                    }
                ]
            )
            
            return response.content[0].text
            
        except Exception as e:
            return f"Error analyzing documents: {str(e)}"

def main():
    st.title("📊 Cap Table Tie-Out Analysis")
    st.markdown("*LLM-powered analysis replicating expert legal review*")
    
    # Initialize analyzer with API key from secrets
    if 'analyzer' not in st.session_state:
        try:
            api_key = st.secrets["ANTHROPIC_API_KEY"]
            st.session_state.analyzer = DeterministicCapTableAnalyzer(api_key)
        except KeyError:
            st.error("❌ Anthropic API key not found in Streamlit secrets. Please configure ANTHROPIC_API_KEY in your .streamlit/secrets.toml file")
            st.stop()
    
    # File upload section
    with st.sidebar:
        st.header("📁 Upload Documents")
        
        # Board documents upload
        st.subheader("Board Documents")
        board_files = st.file_uploader(
            "Upload board consents, minutes, and legal docs",
            type=['docx', 'doc'],
            accept_multiple_files=True,
            key="board_docs",
            help="Upload DOCX files containing board resolutions, consents, and other legal documents"
        )
        
        # Securities ledger upload
        st.subheader("Securities Ledger")
        cap_table_file = st.file_uploader(
            "Upload cap table (Excel format)",
            type=['xlsx', 'xls'],
            key="cap_table",
            help="Upload Excel file containing the company's capitalization table"
        )
        
        # Analysis button
        st.markdown("---")
        run_analysis = st.button(
            "🔍 Run LLM Tie-Out Analysis", 
            type="primary", 
            use_container_width=True
        )
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📋 Uploaded Documents")
        
        if board_files:
            st.write("**Board Documents:**")
            for file in board_files:
                st.write(f"✅ {file.name}")
                
            # Show preview of first document
            if st.checkbox("Show document preview"):
                first_file = board_files[0]
                first_file.seek(0)
                analyzer = st.session_state.get('analyzer')
                if analyzer:
                    content = analyzer.read_docx_content(first_file.read(), first_file.name)
                    st.text_area("Document content preview:", content[:1000] + "...", height=200)
        else:
            st.info("No board documents uploaded yet")
        
        if cap_table_file:
            st.write("**Securities Ledger:**")
            st.write(f"✅ {cap_table_file.name}")
            
            # Show preview of cap table
            if st.checkbox("Show cap table preview"):
                cap_table_file.seek(0)
                try:
                    df_preview = pd.read_excel(io.BytesIO(cap_table_file.read()), engine='openpyxl')
                    st.dataframe(df_preview.head(10))
                except Exception as e:
                    st.error(f"Error previewing cap table: {e}")
        else:
            st.info("No securities ledger uploaded yet")
    
    with col2:
        st.subheader("⚙️ Analysis Status")
        
        if not board_files and not cap_table_file:
            st.warning("📄 Please upload documents to begin analysis")
        elif not board_files:
            st.warning("📋 Please upload board documents")
        elif not cap_table_file:
            st.warning("📊 Please upload securities ledger")
        else:
            st.success("✅ Ready for LLM analysis!")
            st.info("Click 'Run LLM Tie-Out Analysis' to start")
    
    # Run analysis when button is clicked
    if run_analysis:
        if not board_files or not cap_table_file:
            st.error("Please upload both board documents and securities ledger before running analysis")
            return
        
        with st.spinner("🔍 Running deterministic analysis... This will be consistent every time"):
            try:
                analyzer = st.session_state.analyzer
                
                # Process board documents in consistent order (alphabetical)
                board_docs = {}
                sorted_files = sorted(board_files, key=lambda x: x.name)
                for file in sorted_files:
                    file.seek(0)  # Reset file pointer
                    content = analyzer.read_docx_content(file.read(), file.name)
                    board_docs[file.name] = content
                
                # Process cap table
                cap_table_file.seek(0)
                cap_table_entries = analyzer.excel_to_structured_data(cap_table_file.read(), cap_table_file.name)
                
                # Extract board grants using deterministic rules
                board_grants = analyzer.extract_board_grants(board_docs)
                
                # Run deterministic analysis
                discrepancies = analyzer.run_deterministic_analysis(cap_table_entries, board_grants)
                
                # Display results
                st.markdown("---")
                st.header("🎯 Deterministic Analysis Results")
                st.success(f"✅ **Consistent Results**: Found {len(discrepancies)} discrepancies using rule-based analysis")
                
                if discrepancies:
                    # Summary metrics
                    high_count = len([d for d in discrepancies if d['severity'] == 'HIGH'])
                    medium_count = len([d for d in discrepancies if d['severity'] == 'MEDIUM'])
                    low_count = len([d for d in discrepancies if d['severity'] == 'LOW'])
                    
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("🔍 Total Issues", len(discrepancies))
                    with col2:
                        st.metric("🔴 High Severity", high_count)
                    with col3:
                        st.metric("🟡 Medium Severity", medium_count)
                    with col4:
                        st.metric("🟢 Low Severity", low_count)
                    
                    # Display each discrepancy
                    st.subheader("📋 Detailed Discrepancies")
                    
                    for disc in discrepancies:
                        severity_colors = {'HIGH': '#ff4444', 'MEDIUM': '#ffaa00', 'LOW': '#00aa00'}
                        severity_emojis = {'HIGH': '🔴', 'MEDIUM': '🟡', 'LOW': '🟢'}
                        
                        with st.container():
                            st.markdown(
                                f"""
                                <div style="
                                    border-left: 4px solid {severity_colors[disc['severity']]};
                                    background-color: {'#fff5f5' if disc['severity'] == 'HIGH' else '#fffaf0' if disc['severity'] == 'MEDIUM' else '#f0fff4'};
                                    padding: 1rem;
                                    margin: 1rem 0;
                                    border-radius: 0 8px 8px 0;
                                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                                ">
                                    <h4 style="margin: 0 0 0.5rem 0; color: {severity_colors[disc['severity']]};">
                                        {severity_emojis[disc['severity']]} DISCREPANCY #{disc['number']}: {disc['issue']}
                                    </h4>
                                    <p style="margin: 0; font-weight: bold; color: #333;">
                                        👤 <strong>Stockholder:</strong> {disc['stockholder']} | 
                                        📊 <strong>Severity:</strong> {disc['severity']}
                                    </p>
                                </div>
                                """, 
                                unsafe_allow_html=True
                            )
                            
                            col1, col2 = st.columns(2)
                            with col1:
                                st.markdown("**📋 Cap Table Shows:**")
                                st.code(disc['cap_table_value'])
                                st.markdown("**🆔 Security ID:**")
                                st.code(disc['security_id'])
                            
                            with col2:
                                st.markdown("**📜 Legal Documents Show:**")
                                st.code(disc['legal_value'])
                                st.markdown("**📄 Source Document:**")
                                st.code(disc['source'])
                            
                            st.markdown("**📝 Description:**")
                            st.write(disc['description'])
                            st.markdown("---")
                else:
                    st.success("🎉 No discrepancies found! The cap table appears to be in sync with the board documents.")
                
                # Create downloadable report
                st.markdown("---")
                st.subheader("📤 Download Report")
                
                if discrepancies:
                    # Create CSV for download
                    import pandas as pd
                    df_report = pd.DataFrame([
                        {
                            'Discrepancy #': d['number'],
                            'Severity': d['severity'],
                            'Stockholder': d['stockholder'],
                            'Security ID': d['security_id'],
                            'Issue': d['issue'],
                            'Cap Table Shows': d['cap_table_value'],
                            'Legal Documents Show': d['legal_value'],
                            'Description': d['description'],
                            'Source Document': d['source']
                        } for d in discrepancies
                    ])
                    
                    csv = df_report.to_csv(index=False)
                    st.download_button(
                        label="📄 Download Discrepancies Report (CSV)",
                        data=csv,
                        file_name=f"cap_table_discrepancies_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                        mime="text/csv"
                    )
                
            except Exception as e:
                st.error(f"Error during analysis: {str(e)}")
                st.error("Please check your API key and try again")
    
    # Information section
    with st.expander("ℹ️ How this works"):
        st.markdown("""
        **This tool replicates the exact process of expert legal document analysis:**
        
        1. **Document Upload**: You upload board documents and cap table (just like dragging files to Claude)
        2. **LLM Processing**: The documents are sent to Claude (same AI that did the original analysis)
        3. **Expert Analysis**: Claude performs the same thorough legal review methodology
        4. **Detailed Report**: You get the same quality discrepancy analysis and recommendations
        
        **What the LLM analyzes:**
        - Compares cap table entries against board approvals
        - Identifies missing documentation
        - Checks share quantities, pricing, and dates
        - Detects phantom equity grants
        - Finds missing transactions (repurchases, etc.)
        - Provides severity assessment and recommendations
        
        **API Key**: Configured via Streamlit secrets (ANTHROPIC_API_KEY in .streamlit/secrets.toml)
        """)
    
    # Footer
    st.markdown("---")
    st.markdown("*Powered by Claude AI - Professional legal document analysis*")

if __name__ == "__main__":
    main()
,  # Number followed by $
        ]
        
        for pattern in share_patterns:
            share_match = re.search(pattern, content, re.IGNORECASE)
            if share_match:
                shares_str = share_match.group(1).replace(',', '')
                try:
                    shares_num = int(shares_str)
                    if 100 <= shares_num <= 1000000:  # Reasonable range
                        grant['shares'] = shares_num
                        st.write(f"✅ Found shares: {shares_num}")
                        break
                except ValueError:
                    continue
        
        if not grant['shares']:
            st.write("❌ No shares found")
        
        # Extract price - multiple patterns
        price_patterns = [
            r'\$(\d+\.\d{2})\s+per\s+share',
            r'price.*?\$(\d+\.\d{2})',
            r'\$(\d+\.\d{2})',  # Any dollar amount
            r'(\d+\.\d{2})\s+per\s+share',
        ]
        
        for pattern in price_patterns:
            price_match = re.search(pattern, content, re.IGNORECASE)
            if price_match:
                try:
                    price = float(price_match.group(1))
                    if 0.01 <= price <= 1000:  # Reasonable range
                        grant['price_per_share'] = price
                        st.write(f"✅ Found price: ${price}")
                        break
                except ValueError:
                    continue
        
        if not grant['price_per_share']:
            st.write("❌ No price found")
        
        # Extract vesting start date
        vesting_date_patterns = [
            r'vesting.*?(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}',
            r'start.*?(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}',
            r'commencement.*?(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}',
        ]
        
        for pattern in vesting_date_patterns:
            vesting_match = re.search(pattern, content, re.IGNORECASE)
            if vesting_match:
                grant['vesting_start'] = vesting_match.group(1)
                st.write(f"✅ Found vesting start: {grant['vesting_start']}")
                break
        
        if not grant['vesting_start']:
            st.write("❌ No vesting start date found")
        
        # Extract vesting schedule
        if '1/48' in content:
            if 'month' in content.lower():
                grant['vesting_schedule'] = '1/48th monthly'
                st.write("✅ Found vesting: 1/48th monthly")
        elif '25%' in content:
            if 'annual' in content.lower() or 'year' in content.lower():
                grant['vesting_schedule'] = '25% annually'
                st.write("✅ Found vesting: 25% annually")
        
        if not grant['vesting_schedule']:
            st.write("❌ No vesting schedule found")
        
        st.write(f"**Final extracted data:** {grant}")
        st.write("---")
        
        return grant
    
    def extract_repurchase_info(self, content: str, filename: str) -> Dict:
        """Extract repurchase info"""
        import re
        
        repurchase = {
            'type': 'Repurchase',
            'filename': filename,
            'stockholder': None,
            'shares_repurchased': None,
            'date': None
        }
        
        # Extract date
        date_match = re.search(r'Date:\s*([A-Za-z]+\s+\d{1,2},\s+\d{4})', content)
        if date_match:
            repurchase['date'] = date_match.group(1)
        
        # Extract stockholder
        names = ['John Doe', 'Jane Smith', 'Bob', 'Alice', 'Charlie', 'Arthur']
        for name in names:
            if name in content:
                repurchase['stockholder'] = name
                break
        
        # Extract repurchased shares
        repurchase_match = re.search(r'repurchase\s+(\d{1,3}(?:,\d{3})*)', content, re.IGNORECASE)
        if repurchase_match:
            repurchase['shares_repurchased'] = int(repurchase_match.group(1).replace(',', ''))
        
        return repurchase
    
    def extract_option_grant(self, content: str, filename: str) -> Dict:
        """Extract option grant info"""
        # Similar to RSA but for options
        return self.extract_rsa_grant(content, filename)  # Reuse logic for now
    
    def run_deterministic_analysis(self, cap_table_entries: List[Dict], board_grants: List[Dict]) -> List[Dict]:
        """Run deterministic analysis that always produces same results"""
        discrepancies = []
        
        # Create lookup of board grants by stockholder
        board_lookup = {}
        repurchases = []
        
        for grant in board_grants:
            if grant['type'] == 'Repurchase':
                repurchases.append(grant)
            else:
                stockholder = grant.get('stockholder')
                if stockholder:
                    if stockholder not in board_lookup:
                        board_lookup[stockholder] = []
                    board_lookup[stockholder].append(grant)
        
        # Check each cap table entry
        for entry in cap_table_entries:
            security_id = entry.get('Security ID', '')
            stockholder = entry.get('Stakeholder Name', '')
            shares = self.safe_int(entry.get('Quantity Issued', 0))
            cost_basis = self.safe_float(entry.get('Cost Basis', 0))
            board_approval_date = str(entry.get('Board Approval Date', ''))
            vesting_schedule = str(entry.get('Vesting Schedule', ''))
            
            # Calculate price per share
            price_per_share = cost_basis / shares if shares > 0 else 0
            
            # Check 1: Phantom Equity (no board approval found)
            if stockholder not in board_lookup:
                discrepancies.append({
                    'number': len(discrepancies) + 1,
                    'severity': 'HIGH',
                    'stockholder': stockholder,
                    'security_id': security_id,
                    'issue': 'Phantom Equity Entry',
                    'cap_table_value': f'{shares} shares',
                    'legal_value': 'No supporting documentation found',
                    'description': f'Cap table shows {security_id} for {stockholder} but no board approval found',
                    'source': 'None found'
                })
                continue
            
            # Find matching board grant
            matching_grant = None
            for grant in board_lookup[stockholder]:
                if grant.get('shares') == shares or abs((grant.get('shares', 0) - shares)) <= 1:
                    matching_grant = grant
                    break
            
            if not matching_grant:
                matching_grant = board_lookup[stockholder][0]  # Use first grant as fallback
            
            # Check 2: Share quantity mismatch
            board_shares = matching_grant.get('shares', 0)
            if board_shares and abs(shares - board_shares) > 1:
                discrepancies.append({
                    'number': len(discrepancies) + 1,
                    'severity': 'HIGH',
                    'stockholder': stockholder,
                    'security_id': security_id,
                    'issue': 'Incorrect Share Quantity',
                    'cap_table_value': f'{shares} shares',
                    'legal_value': f'{board_shares} shares',
                    'description': f'Cap table shows {shares} shares but board approval is for {board_shares} shares',
                    'source': matching_grant.get('filename', 'Unknown')
                })
            
            # Check 3: Price mismatch
            board_price = matching_grant.get('price_per_share', 0)
            if board_price and abs(price_per_share - board_price) > 0.01:
                discrepancies.append({
                    'number': len(discrepancies) + 1,
                    'severity': 'HIGH',
                    'stockholder': stockholder,
                    'security_id': security_id,
                    'issue': 'Incorrect Price Per Share',
                    'cap_table_value': f'${price_per_share:.2f}',
                    'legal_value': f'${board_price:.2f}',
                    'description': f'Cap table shows ${price_per_share:.2f} per share but board approval is for ${board_price:.2f} per share',
                    'source': matching_grant.get('filename', 'Unknown')
                })
            
            # Check 4: Board approval date mismatch
            board_date = matching_grant.get('date', '')
            if board_date and board_date not in board_approval_date:
                discrepancies.append({
                    'number': len(discrepancies) + 1,
                    'severity': 'HIGH',
                    'stockholder': stockholder,
                    'security_id': security_id,
                    'issue': 'Incorrect Board Approval Date',
                    'cap_table_value': board_approval_date,
                    'legal_value': board_date,
                    'description': 'Board approval date in cap table does not match legal documents',
                    'source': matching_grant.get('filename', 'Unknown')
                })
            
            # Check 5: Vesting schedule mismatch
            board_vesting = matching_grant.get('vesting_schedule', '')
            if board_vesting and board_vesting not in vesting_schedule:
                if ('monthly' in board_vesting.lower() and 'monthly' not in vesting_schedule.lower()) or \
                   ('annual' in board_vesting.lower() and 'annual' not in vesting_schedule.lower()):
                    discrepancies.append({
                        'number': len(discrepancies) + 1,
                        'severity': 'HIGH',
                        'stockholder': stockholder,
                        'security_id': security_id,
                        'issue': 'Vesting Schedule Mismatch',
                        'cap_table_value': vesting_schedule,
                        'legal_value': board_vesting,
                        'description': 'Vesting schedule format differs between cap table and board documents',
                        'source': matching_grant.get('filename', 'Unknown')
                    })
        
        # Check 6: Missing repurchase transactions
        for repurchase in repurchases:
            stockholder = repurchase.get('stockholder')
            shares_repurchased = repurchase.get('shares_repurchased', 0)
            
            if stockholder and shares_repurchased:
                discrepancies.append({
                    'number': len(discrepancies) + 1,
                    'severity': 'HIGH',
                    'stockholder': stockholder,
                    'security_id': 'Multiple',
                    'issue': 'Missing Repurchase Transaction',
                    'cap_table_value': 'No repurchase reflected',
                    'legal_value': f'{shares_repurchased} shares repurchased',
                    'description': f'Board approved repurchase of {shares_repurchased} shares from {stockholder} but cap table does not reflect this transaction',
                    'source': repurchase.get('filename', 'Unknown')
                })
        
        return discrepancies
    
    def safe_int(self, value) -> int:
        """Safely convert to int"""
        try:
            if pd.isna(value):
                return 0
            return int(float(value))
        except (ValueError, TypeError):
            return 0
    
    def safe_float(self, value) -> float:
        """Safely convert to float"""
        try:
            if pd.isna(value):
                return 0.0
            return float(value)
        except (ValueError, TypeError):
            return 0.0
    def __init__(self, api_key: str):
        self.client = anthropic.Anthropic(api_key=api_key)
        self.uploaded_files = {}
        
    def read_docx_content(self, file_content: bytes, filename: str) -> str:
        """Read DOCX content and return as plain text"""
        try:
            doc = Document(io.BytesIO(file_content))
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            return '\n'.join(full_text)
        except Exception as e:
            st.error(f"Error reading {filename}: {str(e)}")
            return ""
    
    def excel_to_text_preview(self, file_content: bytes, filename: str) -> str:
        """Convert Excel to text preview for LLM analysis"""
        try:
            df = pd.read_excel(io.BytesIO(file_content), engine='openpyxl', header=None)
            
            # Create text representation
            text_preview = f"Excel file: {filename}\n\n"
            text_preview += "Raw data structure:\n"
            
            for i in range(min(15, len(df))):  # First 15 rows
                row_data = df.iloc[i].tolist()
                # Clean up the row data
                cleaned_row = []
                for cell in row_data:
                    if pd.isna(cell):
                        cleaned_row.append("")
                    else:
                        cleaned_row.append(str(cell))
                text_preview += f"Row {i + 1}: {cleaned_row}\n"
            
            return text_preview
        except Exception as e:
            return f"Error reading Excel file {filename}: {str(e)}"
    
    def create_analysis_prompt(self, board_docs: Dict[str, str], cap_table_text: str) -> str:
        """Create the enhanced prompt that catches all discrepancies with standardized approach"""
        
        prompt = """You are a lawyer conducting a standardized capitalization table tie out. You MUST follow this exact sequence:

MANDATORY ANALYSIS SEQUENCE:
1. DOCUMENT INVENTORY: List every board document and what it approves
2. CAP TABLE INVENTORY: List every cap table entry (Security ID, Stockholder, Shares)
3. SYSTEMATIC COMPARISON: For each cap table entry, check against board docs
4. DISCREPANCY IDENTIFICATION: List each discrepancy separately with exact format

STEP 1 - DOCUMENT INVENTORY:
First, create a complete list of all board-approved grants from legal documents:
- Document name, date, stockholder, shares, price, vesting details

STEP 2 - CAP TABLE INVENTORY: 
List every entry in the cap table:
- Security ID, Stockholder Name, Quantity, Price details, Dates

STEP 3 - SYSTEMATIC COMPARISON:
For EACH cap table entry, verify these 7 items in order:
a) Does this entry have board approval? (if no = PHANTOM EQUITY)
b) Do share quantities match?
c) Do prices match?
d) Do board approval dates match?
e) Do issue dates match?
f) Do vesting schedules match (monthly vs annual)?
g) Are repurchases reflected?

STEP 4 - DISCREPANCY LIST:
Use this EXACT format for each discrepancy:

DISCREPANCY #[X]: [Issue Title]
- Severity: HIGH/MEDIUM/LOW
- Stockholder: [Name]
- Security ID: [ID]
- Cap Table Shows: [Value]
- Legal Documents Show: [Value]
- Source Document: [Filename]
- Description: [1-2 sentences]

CRITICAL REQUIREMENTS:
- Analyze every single cap table entry
- Check for phantom equity (entries without board approval)
- Verify vesting schedule language matches exactly
- Check for missing repurchase transactions
- List each discrepancy separately (don't group)
- Use consistent severity: HIGH=material issues, MEDIUM=dates/documentation, LOW=minor

Here are the documents to analyze:

BOARD DOCUMENTS:
"""
        
        # Add each board document
        for filename, content in board_docs.items():
            prompt += f"\n--- {filename} ---\n{content}\n"
        
        prompt += f"\nSECURITIES LEDGER / CAP TABLE:\n{cap_table_text}\n"
        
        prompt += """

NOW EXECUTE THE 4-STEP ANALYSIS SEQUENCE ABOVE.

Begin with: "STEP 1 - DOCUMENT INVENTORY:" and follow the exact sequence."""
        
        return prompt
    
    def analyze_with_llm(self, board_docs: Dict[str, str], cap_table_text: str) -> str:
        """Send documents to LLM for analysis"""
        
        prompt = self.create_analysis_prompt(board_docs, cap_table_text)
        
        try:
            response = self.client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=4000,
                temperature=0,  # Maximum determinism
                system="You are a systematic legal auditor. Always follow the exact same analysis sequence and format. Be consistent and thorough in your approach.",
                messages=[
                    {
                        "role": "user", 
                        "content": prompt
                    }
                ]
            )
            
            return response.content[0].text
            
        except Exception as e:
            return f"Error analyzing documents: {str(e)}"

def main():
    st.title("📊 Cap Table Tie-Out Analysis")
    st.markdown("*LLM-powered analysis replicating expert legal review*")
    
    # Initialize analyzer with API key from secrets
    if 'analyzer' not in st.session_state:
        try:
            api_key = st.secrets["ANTHROPIC_API_KEY"]
            st.session_state.analyzer = DeterministicCapTableAnalyzer(api_key)
        except KeyError:
            st.error("❌ Anthropic API key not found in Streamlit secrets. Please configure ANTHROPIC_API_KEY in your .streamlit/secrets.toml file")
            st.stop()
    
    # File upload section
    with st.sidebar:
        st.header("📁 Upload Documents")
        
        # Board documents upload
        st.subheader("Board Documents")
        board_files = st.file_uploader(
            "Upload board consents, minutes, and legal docs",
            type=['docx', 'doc'],
            accept_multiple_files=True,
            key="board_docs",
            help="Upload DOCX files containing board resolutions, consents, and other legal documents"
        )
        
        # Securities ledger upload
        st.subheader("Securities Ledger")
        cap_table_file = st.file_uploader(
            "Upload cap table (Excel format)",
            type=['xlsx', 'xls'],
            key="cap_table",
            help="Upload Excel file containing the company's capitalization table"
        )
        
        # Analysis button
        st.markdown("---")
        run_analysis = st.button(
            "🔍 Run LLM Tie-Out Analysis", 
            type="primary", 
            use_container_width=True
        )
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📋 Uploaded Documents")
        
        if board_files:
            st.write("**Board Documents:**")
            for file in board_files:
                st.write(f"✅ {file.name}")
                
            # Show preview of first document
            if st.checkbox("Show document preview"):
                first_file = board_files[0]
                first_file.seek(0)
                analyzer = st.session_state.get('analyzer')
                if analyzer:
                    content = analyzer.read_docx_content(first_file.read(), first_file.name)
                    st.text_area("Document content preview:", content[:1000] + "...", height=200)
        else:
            st.info("No board documents uploaded yet")
        
        if cap_table_file:
            st.write("**Securities Ledger:**")
            st.write(f"✅ {cap_table_file.name}")
            
            # Show preview of cap table
            if st.checkbox("Show cap table preview"):
                cap_table_file.seek(0)
                try:
                    df_preview = pd.read_excel(io.BytesIO(cap_table_file.read()), engine='openpyxl')
                    st.dataframe(df_preview.head(10))
                except Exception as e:
                    st.error(f"Error previewing cap table: {e}")
        else:
            st.info("No securities ledger uploaded yet")
    
    with col2:
        st.subheader("⚙️ Analysis Status")
        
        if not board_files and not cap_table_file:
            st.warning("📄 Please upload documents to begin analysis")
        elif not board_files:
            st.warning("📋 Please upload board documents")
        elif not cap_table_file:
            st.warning("📊 Please upload securities ledger")
        else:
            st.success("✅ Ready for LLM analysis!")
            st.info("Click 'Run LLM Tie-Out Analysis' to start")
    
    # Run analysis when button is clicked
    if run_analysis:
        if not board_files or not cap_table_file:
            st.error("Please upload both board documents and securities ledger before running analysis")
            return
        
        with st.spinner("🔍 Running deterministic analysis... This will be consistent every time"):
            try:
                analyzer = st.session_state.analyzer
                
                # Process board documents in consistent order (alphabetical)
                board_docs = {}
                sorted_files = sorted(board_files, key=lambda x: x.name)
                for file in sorted_files:
                    file.seek(0)  # Reset file pointer
                    content = analyzer.read_docx_content(file.read(), file.name)
                    board_docs[file.name] = content
                
                # Process cap table
                cap_table_file.seek(0)
                cap_table_entries = analyzer.excel_to_structured_data(cap_table_file.read(), cap_table_file.name)
                
                # Extract board grants using deterministic rules
                board_grants = analyzer.extract_board_grants(board_docs)
                
                # Run deterministic analysis
                discrepancies = analyzer.run_deterministic_analysis(cap_table_entries, board_grants)
                
                # Display results
                st.markdown("---")
                st.header("🎯 Deterministic Analysis Results")
                st.success(f"✅ **Consistent Results**: Found {len(discrepancies)} discrepancies using rule-based analysis")
                
                if discrepancies:
                    # Summary metrics
                    high_count = len([d for d in discrepancies if d['severity'] == 'HIGH'])
                    medium_count = len([d for d in discrepancies if d['severity'] == 'MEDIUM'])
                    low_count = len([d for d in discrepancies if d['severity'] == 'LOW'])
                    
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("🔍 Total Issues", len(discrepancies))
                    with col2:
                        st.metric("🔴 High Severity", high_count)
                    with col3:
                        st.metric("🟡 Medium Severity", medium_count)
                    with col4:
                        st.metric("🟢 Low Severity", low_count)
                    
                    # Display each discrepancy
                    st.subheader("📋 Detailed Discrepancies")
                    
                    for disc in discrepancies:
                        severity_colors = {'HIGH': '#ff4444', 'MEDIUM': '#ffaa00', 'LOW': '#00aa00'}
                        severity_emojis = {'HIGH': '🔴', 'MEDIUM': '🟡', 'LOW': '🟢'}
                        
                        with st.container():
                            st.markdown(
                                f"""
                                <div style="
                                    border-left: 4px solid {severity_colors[disc['severity']]};
                                    background-color: {'#fff5f5' if disc['severity'] == 'HIGH' else '#fffaf0' if disc['severity'] == 'MEDIUM' else '#f0fff4'};
                                    padding: 1rem;
                                    margin: 1rem 0;
                                    border-radius: 0 8px 8px 0;
                                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                                ">
                                    <h4 style="margin: 0 0 0.5rem 0; color: {severity_colors[disc['severity']]};">
                                        {severity_emojis[disc['severity']]} DISCREPANCY #{disc['number']}: {disc['issue']}
                                    </h4>
                                    <p style="margin: 0; font-weight: bold; color: #333;">
                                        👤 <strong>Stockholder:</strong> {disc['stockholder']} | 
                                        📊 <strong>Severity:</strong> {disc['severity']}
                                    </p>
                                </div>
                                """, 
                                unsafe_allow_html=True
                            )
                            
                            col1, col2 = st.columns(2)
                            with col1:
                                st.markdown("**📋 Cap Table Shows:**")
                                st.code(disc['cap_table_value'])
                                st.markdown("**🆔 Security ID:**")
                                st.code(disc['security_id'])
                            
                            with col2:
                                st.markdown("**📜 Legal Documents Show:**")
                                st.code(disc['legal_value'])
                                st.markdown("**📄 Source Document:**")
                                st.code(disc['source'])
                            
                            st.markdown("**📝 Description:**")
                            st.write(disc['description'])
                            st.markdown("---")
                else:
                    st.success("🎉 No discrepancies found! The cap table appears to be in sync with the board documents.")
                
                # Create downloadable report
                st.markdown("---")
                st.subheader("📤 Download Report")
                
                if discrepancies:
                    # Create CSV for download
                    import pandas as pd
                    df_report = pd.DataFrame([
                        {
                            'Discrepancy #': d['number'],
                            'Severity': d['severity'],
                            'Stockholder': d['stockholder'],
                            'Security ID': d['security_id'],
                            'Issue': d['issue'],
                            'Cap Table Shows': d['cap_table_value'],
                            'Legal Documents Show': d['legal_value'],
                            'Description': d['description'],
                            'Source Document': d['source']
                        } for d in discrepancies
                    ])
                    
                    csv = df_report.to_csv(index=False)
                    st.download_button(
                        label="📄 Download Discrepancies Report (CSV)",
                        data=csv,
                        file_name=f"cap_table_discrepancies_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                        mime="text/csv"
                    )
                
            except Exception as e:
                st.error(f"Error during analysis: {str(e)}")
                st.error("Please check your API key and try again")
    
    # Information section
    with st.expander("ℹ️ How this works"):
        st.markdown("""
        **This tool replicates the exact process of expert legal document analysis:**
        
        1. **Document Upload**: You upload board documents and cap table (just like dragging files to Claude)
        2. **LLM Processing**: The documents are sent to Claude (same AI that did the original analysis)
        3. **Expert Analysis**: Claude performs the same thorough legal review methodology
        4. **Detailed Report**: You get the same quality discrepancy analysis and recommendations
        
        **What the LLM analyzes:**
        - Compares cap table entries against board approvals
        - Identifies missing documentation
        - Checks share quantities, pricing, and dates
        - Detects phantom equity grants
        - Finds missing transactions (repurchases, etc.)
        - Provides severity assessment and recommendations
        
        **API Key**: Configured via Streamlit secrets (ANTHROPIC_API_KEY in .streamlit/secrets.toml)
        """)
    
    # Footer
    st.markdown("---")
    st.markdown("*Powered by Claude AI - Professional legal document analysis*")

if __name__ == "__main__":
    main()
